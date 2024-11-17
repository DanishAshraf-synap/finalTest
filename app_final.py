import streamlit as st
from docx import Document
import re
from typing import Dict, List, Optional
import io
import uuid
from embedchain.pipeline import Pipeline as App
from embedchain.config import BaseLlmConfig
import json
import os
remaining_patterns=[]
def initialize_session_state():
    if 'remaining_patterns' not in st.session_state:
        st.session_state.remaining_patterns = []
    if 'patterns' not in st.session_state:
        st.session_state.patterns = []
    if 'current_pattern' not in st.session_state:
        st.session_state.current_pattern = None
    if 'user_inputs' not in st.session_state:
        st.session_state.user_inputs = {}
    if 'messages' not in st.session_state:
        st.session_state.messages = [{
            "role": "assistant",
            "content": "Welcome! Please upload a document to begin."
        }]
    if 'chat_prompt_template' not in st.session_state:
        st.session_state.chat_prompt_template = """
create a question for this item {pattern} for the user with an educated guess from its name and provided context that:
1. Briefly explains what this information is used for in the document.
2. Provides an example of a valid input if applicable.
3. Reminds the user to ensure their input is appropriate and accurate for a legal document.

Format the prompt/question clearly and concisely. Don't start with greetings. Keep it human-like, friendly and short.
"""
    if 'validation_prompt_template' not in st.session_state:
        st.session_state.validation_prompt_template = """
Given the pattern {pattern} and the user input "{user_input}" for a data processing policy document:

Check if the input is empty or just whitespace.
Evaluate if it seems relevant for the given pattern.

Return a dict object with these keys:
"is_valid": boolean (True if valid, False if issues)
"issue": string (user friendly and concise issues, if any)
only and only a dict. never start with python or json etc
"""
def move_to_next_pattern():
    if st.session_state.remaining_patterns:
        # Pop the next pattern
        next_pattern = st.session_state.remaining_patterns.pop(0)
        st.session_state.current_pattern = next_pattern
        # Generate new chat prompt for the next pattern
        prompt = get_chat_response(
            next_pattern,
            st.session_state.chat_prompt_template
        )
        # Append the new prompt
        st.session_state.messages.append({
            "role": "assistant",
            "content": prompt
        })
        return True
    else:
        # No more patterns left, clear current pattern
        st.session_state.current_pattern = None
        return False
def create_embedchain_app(api_key: str):
    if not api_key.strip():
        raise ValueError("API key cannot be empty")
    os.environ["OPENAI_API_KEY"] = api_key
    try:
        return App.from_config(config_path="open.yaml")
    except Exception as e:
        raise Exception(f"Failed to initialize app: {str(e)}")


def extract_patterns_from_docx(file_content: bytes) -> List[str]:
    if not file_content:
        raise ValueError("File content cannot be empty")
    try:
        doc = Document(io.BytesIO(file_content))
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        pattern = r'<([^>]+)>'
        matches = re.findall(pattern, text)
        return list(set(f'<{match}>' for match in matches)),text
    except Exception as e:
        st.error(f"Error processing file: {e}")
        return []
@st.cache_data
def process_docx(content: bytes, user_inputs: Dict[str, str]) -> bytes:
    doc = Document(io.BytesIO(content))
    for para in doc.paragraphs:
        for pattern, replacement in user_inputs.items():
            # print(f"{pattern} -> {replacement}")

            para.text = para.text.replace(f"{pattern}", replacement)
    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream.getvalue()
def get_chat_response(pattern: str, prompt_template: str) -> str:
    prompt = prompt_template.format(pattern=pattern)
    try:
        response = st.session_state.embedchain_app.query(prompt)
        return response
    except Exception as e:
        return f"Error generating question: {str(e)}"

def validate_input(pattern: str, user_input: str) -> tuple[bool, str]:
    validation_prompt = st.session_state.validation_prompt_template.format(
        pattern=pattern,
        user_input=user_input
    )
    
    try:
        # Get the raw response from LLM
        validation_result = st.session_state.embedchain_app.query(validation_prompt)
        
        # Clean the response string - remove any potential leading/trailing whitespace and quotes
        validation_result = validation_result.strip().strip('"\'')
        
        # If the response is already a dict, use it directly
        if isinstance(validation_result, dict):
            result = validation_result
        else:
            # Try to parse the string as JSON
            try:
                result = json.loads(validation_result)
            except json.JSONDecodeError:
                # If JSON parsing fails, try to evaluate as a Python literal
                import ast
                result = ast.literal_eval(validation_result)
        
        # Extract values with proper error handling
        is_valid = bool(result.get('is_valid', False))
        issue = str(result.get('issue', 'Invalid response format'))
        
        return is_valid, issue
        
    except Exception as e:
        st.error(f"Validation error: {str(e)}")
        return False, f"Unable to validate input properly: {str(e)}"

def clear_session_state():
    """Reset all session state variables"""
    st.session_state.remaining_patterns = []
    st.session_state.patterns = []
    st.session_state.current_pattern = None
    st.session_state.user_inputs = {}
    st.session_state.messages = [{
        "role": "assistant",
        "content": "Welcome! Please upload a document to begin."
    }]

def main():
    st.set_page_config(page_title="Complience", layout="wide")
    
    initialize_session_state()

    # Add this near the start of your main() function, before any references to embedchain_app
    if 'embedchain_app' not in st.session_state:
        st.session_state.embedchain_app = None  # or whatever initial value is appropriate

    # Add this near the start of your main() function or before you try to access remaining_patterns
    if 'remaining_patterns' not in st.session_state:
        st.session_state.remaining_patterns = []  # or whatever initial value you want

    # Create two columns layout
    left_panel, main_content = st.columns([2, 5])

    # Left Panel
    with left_panel:
        st.subheader("Settings & Controls")
        
        # API Key input
        api_key = st.text_input(
            "OpenAI API Key", 
            type="password",
            help="Enter your OpenAI API key"
        )
        
        if api_key and not st.session_state.embedchain_app:
            st.session_state.embedchain_app = create_embedchain_app(api_key)

        # Prompt Templates Editor
        with st.expander("Edit Prompt Templates", expanded=False):
            st.text_area(
                "Chat Prompt Template",
                value=st.session_state.chat_prompt_template,
                key="chat_prompt_editor",
                height=200,
                on_change=lambda: setattr(st.session_state, 'chat_prompt_template', st.session_state.chat_prompt_editor)
            )
            
            st.text_area(
                "Validation Prompt Template",
                value=st.session_state.validation_prompt_template,
                key="validation_prompt_editor",
                height=200,
                on_change=lambda: setattr(st.session_state, 'validation_prompt_template', st.session_state.validation_prompt_editor)
            )

        # File Upload
        uploaded_file = st.file_uploader("Upload DOCX file", type="docx")
        if uploaded_file and st.session_state.embedchain_app:
            file_content = uploaded_file.read()
            try :
                st.session_state.embedchain_app.db.reset()
                print("db reset")
            except Exception as e:
                print(e)
            
            patterns,text = extract_patterns_from_docx(file_content)
            st.session_state.embedchain_app.add(text)
            print("text added to db")
            if patterns:
                # Only initialize patterns if they're different from current ones
                if patterns != st.session_state.patterns:
                    st.session_state.patterns = patterns.copy()  # Store all patterns
                    st.session_state.remaining_patterns = patterns.copy()  # Patterns to process
                    st.session_state.current_pattern = None  # Reset current pattern
                    st.session_state.user_inputs = {}  # Reset user inputs
                    st.session_state.original_file = file_content
                    st.success(f"Found {len(patterns)} patterns")
                

            # Display current pattern and navigation controls
            if st.session_state.current_pattern:
                st.info(f"Current Pattern: {st.session_state.current_pattern}")
                
                # Pattern navigation buttons
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("Next Pattern", key="next_pattern"):
                        if move_to_next_pattern():
                            st.rerun()
                        else:
                            # Show completion message
                            st.session_state.messages.append({
                                "role": "assistant",
                                "content": "All patterns have been processed! You can now download your document."
                            })
                            st.rerun()
                
                with col2:
                    if st.button("Retry Current", key="retry_current"):
                        prompt = get_chat_response(
                            st.session_state.current_pattern,
                            st.session_state.chat_prompt_template
                        )
                        st.session_state.messages.append({
                            "role": "assistant",
                            "content": prompt
                        })
                        st.rerun()


        # Add this button in the left panel
        if st.button("Start Over"):
            clear_session_state()
            st.rerun()

    # Main Content - Chat Interface
    with main_content:
        st.markdown("""
            <style>
            .stProgress .st-bo {
                background-color: #f0f2f6;
            }
            .stProgress .st-bp {
                background-color: #00cc00;
            }
            </style>
            """, unsafe_allow_html=True)
        
        st.markdown("### Document Pattern Processor")
        
        # Display chat messages
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

        # Initialize first pattern if none is selected
        if (st.session_state.embedchain_app and 
            not st.session_state.current_pattern and 
            st.session_state.remaining_patterns):
            st.session_state.current_pattern = st.session_state.remaining_patterns.pop(0)
            prompt = get_chat_response(
                st.session_state.current_pattern,
                st.session_state.chat_prompt_template
            )
            st.session_state.messages.append({"role": "assistant", "content": prompt})
            st.rerun()

        # Chat input
        if st.session_state.current_pattern:
            user_input = st.chat_input("Your response")
            if user_input:
                st.session_state.messages.append({"role": "user", "content": user_input})
                is_valid, issue = validate_input(st.session_state.current_pattern, user_input)
                
                if is_valid:
                    st.session_state.user_inputs[st.session_state.current_pattern] = user_input
                    if st.session_state.remaining_patterns:
                        st.session_state.messages.append({
                            "role": "assistant", 
                            "content": "Thank you! Click 'Next Pattern' to continue."
                        })
                    else:
                        # No more patterns left, show download button
                        processed_content = process_docx(
                            st.session_state.original_file,
                            st.session_state.user_inputs
                        )
                        st.download_button(
                            label="Download Processed Document",
                            data=processed_content,
                            file_name=f"processed_{uuid.uuid4()}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        st.session_state.messages.append({
                            "role": "assistant", 
                            "content": "All patterns have been processed! You can now download your document."
                        })
                else:
                    st.session_state.messages.append({
                        "role": "assistant", 
                        "content": f"There seems to be an issue with your input: {issue}. Please try again."
                    })
                st.rerun()
        # Move the download button section outside the chat input condition
        if not st.session_state.remaining_patterns and st.session_state.user_inputs and st.session_state.original_file:
            processed_content = process_docx(
                st.session_state.original_file,
                st.session_state.user_inputs
            )
            st.download_button(
                label="Download Processed Document",
                data=processed_content,
                file_name=f"processed_{uuid.uuid4()}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        if st.session_state.patterns:
            total_patterns = len(st.session_state.patterns)
            completed_patterns = len(st.session_state.user_inputs)
            st.progress(completed_patterns / total_patterns)
            st.write(f"Progress: {completed_patterns}/{total_patterns} patterns processed")

        if st.session_state.user_inputs:
            with st.expander("View Current Inputs"):
                for pattern, value in st.session_state.user_inputs.items():
                    st.write(f"{pattern}: {value}")

if __name__ == "__main__":
    main()